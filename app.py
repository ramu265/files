import os
import re
from io import BytesIO
from docx import Document
from flask import (
    Flask,
    flash,
    redirect,
    render_template,
    request,
    send_file,
)

app = Flask(__name__)
app.secret_key = "legal_cleaner_secret_key"


# -----------------------------
# TEXT CLEANING FUNCTION
# -----------------------------
def clean_text_data(text):

    # Remove unwanted slash symbols
    cleaned = re.sub(r"/", "", text)

    # Remove page headers like:
    # Page 25 of [2025] KEHC...
    cleaned = re.sub(
        r"Page\s+\d+\s+of\s+.*?(?=\n|$)",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )

    # Remove extra spaces/tabs
    cleaned = re.sub(r"[ \t]+", " ", cleaned)

    # Remove multiple blank lines
    cleaned = re.sub(r"\n\s*\n+", "\n\n", cleaned)

    # Clean line by line
    lines = cleaned.splitlines()

    final_lines = []
    for line in lines:
        line = line.strip()

        if line:
            final_lines.append(line)

    return "\n".join(final_lines)


# -----------------------------
# HOME ROUTE
# -----------------------------
@app.route("/", methods=["GET", "POST"])
def index():

    if request.method == "POST":

        # Check file exists
        if "file" not in request.files:
            flash("No file uploaded")
            return redirect(request.url)

        file = request.files["file"]

        # Empty filename
        if file.filename == "":
            flash("Please select a file")
            return redirect(request.url)

        filename = file.filename
        extension = os.path.splitext(filename)[1].lower()

        try:

            # ------------------------------------------------
            # DOCX / DOC FILE PROCESSING
            # ------------------------------------------------
            if extension in [".docx", ".doc"]:

                doc = Document(file)

                full_text = []

                for para in doc.paragraphs:
                    full_text.append(para.text)

                raw_content = "\n".join(full_text)

                # Clean text
                cleaned_content = clean_text_data(raw_content)

                # Create output document
                output_doc = Document()

                for line in cleaned_content.split("\n"):
                    output_doc.add_paragraph(line)

                # Save to memory
                memory_file = BytesIO()
                output_doc.save(memory_file)
                memory_file.seek(0)

                # Output filename
                if extension == ".docx":
                    new_filename = f"Cleaned_{filename}"
                else:
                    new_filename = f"Cleaned_{filename}.docx"

                return send_file(
                    memory_file,
                    as_attachment=True,
                    download_name=new_filename,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            # ------------------------------------------------
            # TXT FILE PROCESSING
            # ------------------------------------------------
            elif extension == ".txt":

                raw_content = file.read().decode(
                    "utf-8",
                    errors="ignore"
                )

                cleaned_content = clean_text_data(raw_content)

                memory_file = BytesIO()
                memory_file.write(cleaned_content.encode("utf-8"))
                memory_file.seek(0)

                return send_file(
                    memory_file,
                    as_attachment=True,
                    download_name=f"Cleaned_{filename}",
                    mimetype="text/plain",
                )

            # ------------------------------------------------
            # UNSUPPORTED FORMAT
            # ------------------------------------------------
            else:
                flash(
                    "Unsupported file format! Upload .docx, .doc or .txt files only."
                )
                return redirect(request.url)

        except Exception as e:
            flash(f"Error processing file: {str(e)}")
            return redirect(request.url)

    return render_template("index.html")


# -----------------------------
# RUN APP
# -----------------------------
if __name__ == "__main__":
    app.run(debug=True)